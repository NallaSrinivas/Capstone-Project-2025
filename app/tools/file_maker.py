import os
from datetime import datetime

from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle


def save_file(topic: str, content: str, name_file: str, file_format: str = "md") -> dict:
    """
    Saves a summary to /outputs as Markdown, PDF, or DOCX.
    Designed for OpenAI tool calling (returns dict).
    """
    # Always save to project_root/outputs
    root_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../"))
    output_dir = os.path.join(root_dir, "outputs")
    os.makedirs(output_dir, exist_ok=True)

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")

    name_file = name_file.strip()
    file_format = file_format.lower().strip()

    try:
        if file_format == "docx":
            return _save_docx(output_dir, name_file, topic, content, timestamp)
        elif file_format == "pdf":
            return _save_pdf(output_dir, name_file, topic, content, timestamp)
        else:
            return _save_markdown(output_dir, name_file, topic, content, timestamp)

    except Exception as e:
        return {"success": False, "message": str(e)}




# ============================================================
# MARKDOWN
# ============================================================

def _save_markdown(output_dir, name_file, topic, content, timestamp):
    filename = f"{name_file}.md"
    filepath = os.path.join(output_dir, filename)

    md = f"# {topic}\n\n{content}\n\n---\nSaved on: {timestamp}"

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(md)

    return {"success": True, "path": filepath}



# ============================================================
# DOCX
# ============================================================

def _save_docx(output_dir, name_file, topic, content, timestamp):
    filename = f"{name_file}.docx"
    filepath = os.path.join(output_dir, filename)

    doc = Document()
    doc.add_heading(topic, level=1)

    for line in content.split("\n"):
        doc.add_paragraph(line)

    doc.add_paragraph(f"\nSaved on: {timestamp}")

    doc.save(filepath)
    return {"success": True, "path": filepath}




# ============================================================
# PDF  (FIXED: MULTI-PAGE, FULL CONTENT)
# ============================================================

def _save_pdf(output_dir, name_file, topic, content, timestamp):
    filename = f"{name_file}.pdf"
    filepath = os.path.join(output_dir, filename)

    styles = getSampleStyleSheet()

    # Create flowing text style
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        leading=16,
        spaceAfter=10
    )

    pdf = SimpleDocTemplate(filepath)

    story = []

    # Title
    story.append(Paragraph(f"<b>{topic}</b>", styles["Title"]))
    story.append(Paragraph("<br/>", body_style))

    # Split into paragraphs so ReportLab can auto-page-break
    for line in content.split("\n"):
        if line.strip():
            story.append(Paragraph(line, body_style))
        else:
            # empty line
            story.append(Paragraph("<br/>", body_style))

    # Footer on new page
    story.append(PageBreak())
    story.append(Paragraph(f"Saved on: {timestamp}", styles["Italic"]))

    pdf.build(story)

    return {"success": True, "path": filepath}
